"""Извлечение текста из загруженных резюме (PDF / DOCX / TXT)."""

import io

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class ExtractionError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith((".txt", ".md")):
        return _extract_plain(data)
    raise ExtractionError(
        f"Неподдерживаемый формат файла: {filename}. Поддерживаются PDF, DOCX, TXT."
    )


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ExtractionError(f"Не удалось прочитать PDF: {exc}") from exc
    text = "\n".join(pages).strip()
    if text:
        return text

    # Текстового слоя нет — вероятно, скан. Пробуем OCR.
    from ..config import get_settings

    if get_settings().ocr_enabled:
        ocr_text = _ocr_pdf(data)
        if ocr_text:
            return ocr_text

    raise ExtractionError(
        "PDF не содержит извлекаемого текста. Для сканов нужен установленный "
        "OCR (tesseract + poppler) либо загрузите текстовый PDF/DOCX."
    )


def _ocr_pdf(data: bytes) -> str:
    """OCR сканированного PDF. Возвращает '' если OCR недоступен или ничего не распознал."""
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        return ""

    from ..config import get_settings

    settings = get_settings()
    try:
        images = convert_from_bytes(
            data, dpi=200, last_page=settings.ocr_max_pages
        )
    except Exception:
        # poppler не установлен или файл битый
        return ""

    parts: list[str] = []
    for image in images:
        try:
            parts.append(pytesseract.image_to_string(image, lang=settings.ocr_lang))
        except Exception:
            try:
                parts.append(pytesseract.image_to_string(image))
            except Exception:
                return ""  # tesseract-бинарь недоступен
    return "\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    import docx

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Не удалось прочитать DOCX: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(part for part in parts if part.strip()).strip()
    if not text:
        raise ExtractionError("DOCX не содержит текста")
    return text


def _extract_plain(data: bytes) -> str:
    for encoding in ("utf-8", "cp1251"):
        try:
            text = data.decode(encoding).strip()
            if text:
                return text
        except UnicodeDecodeError:
            continue
    raise ExtractionError("Не удалось декодировать текстовый файл")
